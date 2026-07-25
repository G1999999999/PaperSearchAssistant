"""
文档加载与预处理工具。

这里主要做两件事：
1. 将原始输入（字符串或简单文件）加载为纯文本。
2. 在一个地方集中做“轻量清洗”，方便在面试时讲输入预处理策略。
"""

from __future__ import annotations

from pathlib import Path
from typing import Tuple
from uuid import uuid4

from bs4 import BeautifulSoup
from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader


def _basic_clean(text: str) -> str:
    """对文本做一些非常轻量的清洗。

    - 去掉首尾空白
    - 统一换行风格
    - 合并连续空行
    """

    normalized = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    # 合并多余空行
    lines = [line.rstrip() for line in normalized.split("\n")]
    cleaned_lines: list[str] = []
    blank = False
    for line in lines:
        if line == "":
            if not blank:
                cleaned_lines.append(line)
            blank = True
        else:
            cleaned_lines.append(line)
            blank = False
    return "\n".join(cleaned_lines)


def load_plain_text(text: str, source: str = "inline") -> Tuple[str, dict]:
    """从字符串加载文本，并返回清洗后的内容和 metadata。"""

    cleaned = _basic_clean(text)
    parent_id = uuid4().hex
    metadata = {"source": source, "parent_id": parent_id}
    return cleaned, metadata


def _metadata_from_path(p: Path, parent_id: str | None = None) -> dict:
    pid = parent_id or str(p)
    return {"source": str(p), "name": p.name, "parent_id": pid}


def _load_txt(p: Path) -> Tuple[str, dict]:
    content = p.read_text(encoding="utf-8")
    cleaned = _basic_clean(content)
    return cleaned, _metadata_from_path(p)


def _load_md(p: Path) -> Tuple[str, dict]:
    content = p.read_text(encoding="utf-8")
    cleaned = _basic_clean(content)
    return cleaned, _metadata_from_path(p)


def _load_html(p: Path) -> Tuple[str, dict]:
    html = p.read_text(encoding="utf-8")
    soup = BeautifulSoup(html, "html.parser")
    text = soup.get_text("\n")
    cleaned = _basic_clean(text)
    return cleaned, _metadata_from_path(p)


def _load_docx(p: Path) -> Tuple[str, dict]:
    doc = DocxDocument(str(p))
    paras = [para.text for para in doc.paragraphs if para.text]
    text = "\n".join(paras)
    cleaned = _basic_clean(text)
    return cleaned, _metadata_from_path(p)


def _load_excel(p: Path) -> Tuple[str, dict]:
    wb = load_workbook(str(p), data_only=True)
    lines: list[str] = []
    for sheet in wb.worksheets:
        lines.append(f"# 工作表：{sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                lines.append("\t".join(cells))
        lines.append("")
    text = "\n".join(lines)
    cleaned = _basic_clean(text)
    return cleaned, _metadata_from_path(p)


def load_file(path: str) -> Tuple[str, dict]:
    """从文件加载文本，支持 txt/md/html/pdf/docx/xlsx 等常见格式。"""

    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(path)

    suffix = p.suffix.lower()
    if suffix in {".txt"}:
        return _load_txt(p)
    if suffix in {".md", ".markdown"}:
        return _load_md(p)
    if suffix in {".html", ".htm"}:
        return _load_html(p)
    if suffix in {".pdf"}:
        return load_pdf(str(p))
    if suffix in {".docx"}:
        return _load_docx(p)
    if suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        return _load_excel(p)

    # 默认按纯文本处理
    return _load_txt(p)


def load_pdf(path: str, parent_id: str | None = None) -> Tuple[str, dict]:
    """从 PDF 文件加载文本，返回清洗后的内容和 metadata。"""

    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(path)

    reader = PdfReader(str(p))
    pages_text: list[str] = []
    for page in reader.pages:
        try:
            pages_text.append(page.extract_text() or "")
        except Exception:
            continue
    content = "\n".join(pages_text)
    cleaned = _basic_clean(content)
    pid = parent_id or p.stem
    metadata = {"source": str(p), "name": p.name, "parent_id": pid}
    return cleaned, metadata


def load_pdf_chunks_pymupdf(
    path: str,
    parent_id: str | None = None,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
) -> list[LCDocument]:
    """使用 PyMuPDFLoader 按页读取 PDF，并切片为可直接入向量库的 chunks。"""

    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(path)

    try:
        from langchain_community.document_loaders import PyMuPDFLoader
    except ImportError as exc:
        raise RuntimeError(
            "请安装 langchain-community 与 pymupdf: pip install langchain-community pymupdf"
        ) from exc

    loader = PyMuPDFLoader(str(p))
    page_docs = loader.load()

    pid = parent_id or p.stem
    normalized_docs: list[LCDocument] = []
    for d in page_docs:
        content = _basic_clean(d.page_content or "")
        if not content:
            continue
        md = dict(d.metadata or {})
        md.update(
            {
                "source": str(p),
                "name": p.name,
                "parent_id": pid,
            }
        )
        normalized_docs.append(LCDocument(page_content=content, metadata=md))

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )
    chunks = splitter.split_documents(normalized_docs)
    for i, chunk in enumerate(chunks):
        chunk.metadata = dict(chunk.metadata or {})
        chunk.metadata["chunk_index"] = i
    return chunks


def prepare_document(text: str, source: str = "inline") -> Tuple[str, dict]:
    """统一封装，返回 (content, metadata) 结构，方便上层调用。"""

    return load_plain_text(text=text, source=source)

